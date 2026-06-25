from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import io

# Word生成（修复临时文件二进制读取问题）
def create_exam_word(questions_text: str, title: str, has_answer: bool) -> bytes:
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(11)
    # 标题
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = 1
    mode_text = "（含答案及详细解析）" if has_answer else "（无答案版）"
    doc.add_paragraph(mode_text).alignment = 1
    doc.add_paragraph("")
    lines = questions_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if not has_answer and ("【答案】" in line or "【解析】" in line):
            continue
        doc.add_paragraph(line)
    # 内存字节流输出，不依赖本地临时文件，彻底修复下载失效
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# PDF试卷导出
def create_exam_pdf(questions_text: str, title: str, has_answer: bool) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont("SimSun", 16)
    c.drawCentredString(width/2, y, title)
    y -= 25
    c.setFont("SimSun", 12)
    mode_text = "（含答案及详细解析）" if has_answer else "（无答案版）"
    c.drawCentredString(width/2, y, mode_text)
    y -= 40
    lines = questions_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if not has_answer and ("【答案】" in line or "【解析】" in line):
            continue
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(30, y, line)
        y -= 20
    c.save()
    buf.seek(0)
    return buf.getvalue()

# Excel题库批量导出
def export_questions_excel(question_list) -> bytes:
    data = []
    for q in question_list:
        data.append({
            "题目ID": q["id"],
            "试卷标题": q["title"],
            "使用场景": q["scene"],
            "器官系统": ",".join(q["organ_systems"]),
            "考核主题": ",".join(q["topics"]),
            "难度等级": ",".join(q["diff_levels"]),
            "题型": ",".join(q["q_types"]),
            "完整试题内容": q["full_content"],
            "创建时间": q["create_time"]
        })
    df = pd.DataFrame(data)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="药学题库")
    buf.seek(0)
    return buf.getvalue()