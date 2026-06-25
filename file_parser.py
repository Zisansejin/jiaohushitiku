import os
import markdown
from pypdf import PdfReader
from docx import Document
import openpyxl

def parse_file(file_like) -> str:
    """
    兼容两种入参：
    1. streamlit UploadedFile 内存文件流
    2. 本地文件路径字符串
    """
    # 判断入参是文件路径还是内存流
    if isinstance(file_like, str):
        file_path = file_like
        ext = os.path.splitext(file_path)[1].lower()
    else:
        # UploadedFile 对象，读取后缀
        file_path = None
        ext = os.path.splitext(file_like.name)[1].lower()

    content = ""
    try:
        if ext == ".docx":
            if file_path:
                doc = Document(file_path)
            else:
                doc = Document(file_like)
            for para in doc.paragraphs:
                content += para.text + "\n"

        elif ext == ".pdf":
            # 直接传入内存文件流，不落地本地文件
            reader = PdfReader(file_like)
            for page in reader.pages:
                # 开启layout布局提取，大幅改善中文PDF提取
                page_text = page.extract_text(extraction_mode="layout") or ""
                content += page_text + "\n"

        elif ext in [".xlsx", ".xls"]:
            if file_path:
                wb = openpyxl.load_workbook(file_path)
            else:
                wb = openpyxl.load_workbook(file_like)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_str = " | ".join([str(i) for i in row if i is not None])
                    content += row_str + "\n"

        elif ext == ".md":
            if file_path:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = file_like.getvalue().decode("utf-8")

        elif ext == ".txt":
            if file_path:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = file_like.getvalue().decode("utf-8")
    except Exception as e:
        content = f"【文件解析异常】{str(e)}\n原始内容读取失败，请检查文件是否为可复制文本PDF，非扫描图片版"
    return content.strip()