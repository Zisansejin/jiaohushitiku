import sys
import os
import tempfile
import shutil
from streamlit.web import cli as stcli

# ====================== 内嵌所有依赖模块代码 ======================
# 1. api_client.py 完整代码
class MultiAIClient:
    def __init__(self, model_type: str, api_key: str):
        self.model_type = model_type
        self.api_key = api_key
        self.base_url = ""
        self.model_name = ""
        self.headers = {"Content-Type": "application/json"}
        self.headers["Authorization"] = f"Bearer {self.api_key}"
        if model_type == "deepseek":
            self.base_url = "https://api.deepseek.com/v1/chat/completions"
            self.model_name = "deepseek-chat"
        elif model_type == "openai":
            self.base_url = "https://api.openai.com/v1/chat/completions"
            self.model_name = "gpt-3.5-turbo"
        elif model_type == "tongyi":
            self.base_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
            self.model_name = "qwen-turbo"
        elif model_type == "ernie":
            self.base_url = "https://qianfan.baidubce.com/v2/chat/completions"
            self.model_name = "ernie-3.5"

    def chat_completion(self, prompt: str, temperature=0.3):
        import requests
        sys_prompt = "你是资深临床药学教学专家，严格按照药学规培大纲、处方审核规范出题，答案解析专业严谨，分级清晰，格式规整。"
        payload = {
            "model": self.model_name,
            "messages": [
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": prompt}
            ],
            "temperature": temperature,
            "max_tokens": 4096
        }
        resp = requests.post(self.base_url, headers=self.headers, json=payload, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"]

# 2. file_parser.py 完整代码
def parse_file(file_like) -> str:
    import os
    from pypdf import PdfReader
    from docx import Document
    import openpyxl
    if isinstance(file_like, str):
        file_path = file_like
        ext = os.path.splitext(file_path)[1].lower()
    else:
        file_path = None
        ext = os.path.splitext(file_like.name)[1].lower()
    content = ""
    try:
        if ext == ".docx":
            if file_path:
                doc = Document(file_path)
            else:
                doc = Document(file_like)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif ext == ".pdf":
            reader = PdfReader(file_like)
            for page in reader.pages:
                page_text = page.extract_text(extraction_mode="layout") or ""
                content += page_text + "\n"
        elif ext in [".xlsx", ".xls"]:
            if file_path:
                wb = openpyxl.load_workbook(file_path)
            else:
                wb = openpyxl.load_workbook(file_like)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_str = " | ".join([str(i) for i in row if i is not None])
                    content += row_str + "\n"
        elif ext == ".md":
            if file_path:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = file_like.getvalue().decode("utf-8")
        elif ext == ".txt":
            if file_path:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = file_like.getvalue().decode("utf-8")
    except Exception as e:
        content = f"【文件解析异常】{str(e)}\n原始内容读取失败，请检查文件是否为可复制文本PDF，非扫描图片版"
    return content.strip()

# 3. doc_export.py 完整代码
def create_exam_word(questions_text: str, title: str, has_answer: bool) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    import io
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = 1
    mode_text = "（含答案及详细解析）" if has_answer else "（无答案版）"
    doc.add_paragraph(mode_text).alignment = 1
    doc.add_paragraph("")
    lines = questions_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if not has_answer and ("【答案】" in line or "【解析】" in line):
            continue
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def create_exam_pdf(questions_text: str, title: str, has_answer: bool) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    import io
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont("SimSun", 16)
    c.drawCentredString(width/2, y, title)
    y -= 25
    c.setFont("SimSun", 12)
    mode_text = "（含答案及详细解析）" if has_answer else "（无答案版）"
    c.drawCentredString(width/2, y, mode_text)
    y -= 40
    lines = questions_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if not has_answer and ("【答案】" in line or "【解析】" in line):
            continue
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(30, y, line)
        y -= 20
    c.save()
    buf.seek(0)
    return buf.getvalue()

def export_questions_excel(question_list) -> bytes:
    import pandas as pd
    import io
    data = []
    for q in question_list:
        data.append({
            "题目ID": q["id"],
            "试卷标题": q["title"],
            "使用场景": q["scene"],
            "器官系统": ",".join(q["organ_systems"]),
            "考核主题": ",".join(q["topics"]),
            "难度等级": ",".join(q["diff_levels"]),
            "题型": ",".join(q["q_types"]),
            "完整试题内容": q["full_content"],
            "创建时间": q["create_time"]
        })
    df = pd.DataFrame(data)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="药学题库")
    buf.seek(0)
    return buf.getvalue()

# 4. db.py 完整代码
import sqlite3
import json
import datetime
DB_FILE = "pharm_db.sqlite"
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute('''
    CREATE TABLE IF NOT EXISTS question_bank (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        full_content TEXT,
        scene TEXT,
        organ_systems TEXT,
        topics TEXT,
        diff_levels TEXT,
        q_types TEXT,
        master_levels TEXT,
        create_time DATETIME
    )
    ''')
    cur.execute('''
    CREATE TABLE IF NOT EXISTS drug_list (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        drug_name TEXT UNIQUE,
        organ_system TEXT,
        drug_class TEXT
    )
    ''')
    cur.execute('''
    CREATE TABLE IF NOT EXISTS prompt_template (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        template_name TEXT UNIQUE,
        prompt_content TEXT,
        model_type TEXT,
        create_time DATETIME
    )
    ''')
    conn.commit()
    conn.close()
def add_question_to_db(data: dict):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur.execute('''
    INSERT INTO question_bank
    (title, full_content, scene, organ_systems, topics, diff_levels, q_types, master_levels, create_time)
    VALUES (?,?,?,?,?,?,?,?,?)
    ''', (
        data["title"],
        data["full_content"],
        data["scene"],
        json.dumps(data["organ_systems"], ensure_ascii=False),
        json.dumps(data["topics"], ensure_ascii=False),
        json.dumps(data["diff_levels"], ensure_ascii=False),
        json.dumps(data["q_types"], ensure_ascii=False),
        json.dumps(data["master_levels"], ensure_ascii=False),
        now
    ))
    conn.commit()
    conn.close()
def query_questions(scene=None, organ_system=None, topic=None, diff=None, q_type=None):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    sql = "SELECT * FROM question_bank WHERE 1=1"
    params = []
    if scene:
        sql += " AND scene=?"
        params.append(scene)
    if organ_system:
        sql += " AND organ_systems LIKE ?"
        params.append(f'%{organ_system}%')
    if topic:
        sql += " AND topics LIKE ?"
        params.append(f'%{topic}%')
    if diff:
        sql += " AND diff_levels LIKE ?"
        params.append(f'%{diff}%')
    if q_type:
        sql += " AND q_types LIKE ?"
        params.append(f'%{q_type}%')
    sql += " ORDER BY create_time DESC"
    cur.execute(sql, params)
    rows = cur.fetchall()
    res = []
    for r in rows:
        res.append({
            "id": r[0],
            "title": r[1],
            "full_content": r[2],
            "scene": r[3],
            "organ_systems": json.loads(r[4]),
            "topics": json.loads(r[5]),
            "diff_levels": json.loads(r[6]),
            "q_types": json.loads(r[7]),
            "master_levels": json.loads(r[8]),
            "create_time": r[9]
        })
    conn.close()
    return res
def get_question_by_ids(id_list):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    placeholders = ",".join(["?"] * len(id_list))
    cur.execute(f"SELECT * FROM question_bank WHERE id IN ({placeholders})", id_list)
    rows = cur.fetchall()
    res = []
    for r in rows:
        res.append({
            "id": r[0],
            "title": r[1],
            "full_content": r[2],
            "scene": r[3],
            "organ_systems": json.loads(r[4]),
            "topics": json.loads(r[5]),
            "diff_levels": json.loads(r[6]),
            "q_types": json.loads(r[7]),
            "master_levels": json.loads(r[8]),
            "create_time": r[9]
        })
    conn.close()
    return res
def add_drug(drug_name: str, organ_system: str, drug_class: str):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("INSERT OR IGNORE INTO drug_list (drug_name, organ_system, drug_class) VALUES (?,?,?)",
                (drug_name, organ_system, drug_class))
    conn.commit()
    conn.close()
def get_all_drugs():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("SELECT * FROM drug_list")
    rows = cur.fetchall()
    res = []
    for r in rows:
        res.append({"id": r[0], "drug_name": r[1], "organ_system": r[2], "drug_class": r[3]})
    conn.close()
    return res
def save_prompt_template(tpl_name: str, content: str, model_type: str):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur.execute('''
    INSERT OR REPLACE INTO prompt_template (template_name, prompt_content, model_type, create_time)
    VALUES (?,?,?,?)
    ''', (tpl_name, content, model_type, now))
    conn.commit()
    conn.close()
def get_all_templates():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("SELECT * FROM prompt_template ORDER BY create_time DESC")
    rows = cur.fetchall()
    res = []
    for r in rows:
        res.append({
            "id": r[0],
            "template_name": r[1],
            "prompt_content": r[2],
            "model_type": r[3],
            "create_time": r[4]
        })
    conn.close()
    return res
init_db()
base_drugs = [
    ("阿莫西林", "抗感染系统", "青霉素类"),
    ("头孢曲松", "抗感染系统", "头孢三代"),
    ("氯沙坦", "心血管系统", "ARB降压药"),
    ("二甲双胍", "内分泌/糖尿病", "双胍类降糖药"),
    ("阿托伐他汀", "心血管系统", "他汀降脂"),
    ("呋塞米", "肾脏泌尿", "袢利尿剂")
]
for dname, sys, cls in base_drugs:
    add_drug(dname, sys, cls)

# ====================== 原main.py完整界面代码 ======================
def run_app():
    import streamlit as st
    import io
    st.set_page_config(page_title="药学规培AI题库生成系统", layout="wide")
    if "ai_client" not in st.session_state:
        st.session_state.ai_client = None
    if "outline_content" not in st.session_state:
        st.session_state.outline_content = ""
    if "parsed_outline_text" not in st.session_state:
        st.session_state.parsed_outline_text = ""
    if "generated_questions" not in st.session_state:
        st.session_state.generated_questions = ""
    if "selected_q_ids" not in st.session_state:
        st.session_state.selected_q_ids = []
    with st.sidebar:
        st.header("⚙️ 全局AI模型配置")
        model_type = st.selectbox("选择大模型", ["deepseek", "openai", "tongyi", "ernie"], format_func=lambda x: {
            "deepseek":"DeepSeek", "openai":"OpenAI", "tongyi":"通义千问", "ernie":"文心一言"
        }[x])
        api_key = st.text_input("API Key", type="password")
        if api_key:
            st.session_state.ai_client = MultiAIClient(model_type, api_key)
        st.divider()
        st.subheader("药物库管理")
        drug_list = get_all_drugs()
        drug_names = [d["drug_name"] for d in drug_list]
        new_drug = st.text_input("新增药物名称")
        new_sys = st.selectbox("所属系统", ["抗感染系统","心血管系统","内分泌/糖尿病","肾脏泌尿","消化胃肠"])
        new_cls = st.text_input("药物分类")
        if st.button("添加药物到库") and new_drug:
            add_drug(new_drug, new_sys, new_cls)
            st.rerun()
        st.divider()
        st.subheader("Prompt模板管理")
        all_tpl = get_all_templates()
        tpl_names = [t["template_name"] for t in all_tpl]
        sel_tpl = st.selectbox("加载保存的模板", ["无"] + tpl_names)
    tab1, tab2, tab3 = st.tabs(["🤖 AI出题生成", "📚 题库筛选管理", "✍️ 在线自测答题"])
    with tab1:
        col_left, col_right = st.columns([1,1])
        with col_left:
            st.subheader("出题参数配置")
            upload_file = st.file_uploader("导入规培大纲 Word/PDF/Excel/MD/TXT", type=["docx","pdf","xlsx","md","txt"])
            if upload_file is not None:
                with st.spinner("正在解析文档内容..."):
                    parse_result = parse_file(upload_file)
                    st.session_state.parsed_outline_text = parse_result
                    st.session_state.outline_content = parse_result
                    st.success(f"✅ {upload_file.name} 解析完成！已加载到下方编辑框")
                    st.info(f"解析内容预览：{parse_result[:200]}")
            user_edit_outline = st.text_area(
                "大纲内容编辑（可手动修改补充）",
                value=st.session_state.outline_content,
                height=120
            )
            st.session_state.outline_content = user_edit_outline
            st.divider()
            master_level = st.multiselect("大纲掌握等级", ["了解","熟悉","掌握"], default=["掌握"])
            diff_raw = st.multiselect("认知难度", ["L1 简单","L2 一般","L3 困难"], default=["L1 简单","L2 一般","L3 困难"])
            diff_map = {"L1 简单":"L1","L2 一般":"L2","L3 困难":"L3"}
            diff_list = [diff_map[i] for i in diff_raw]
            scene = st.selectbox("出题场景", ["门诊处方","住院医嘱","药历点评","用药教育","本科期末考试","规培考核"])
            organ_sys = st.multiselect("考核器官系统", ["抗感染系统","心血管系统","内分泌/糖尿病","肾脏泌尿","消化胃肠","神经精神"], default=["抗感染系统","内分泌/糖尿病"])
            select_drugs = st.multiselect("指定药物出题（可选）", drug_names)
            topic_list = st.multiselect("考核主题", ["药物相互作用","剂型/用法用量","老年人用药禁忌","重复用药","肝肾功能剂量调整","抗菌不合理使用"], default=["抗菌不合理使用","药物相互作用"])
            q_type = st.multiselect("题型", ["单选题","多选题","判断题","填空题","简答题"], default=["单选题","多选题"])
            q_count = st.number_input("生成题量", min_value=1, max_value=50, value=5)
            st.divider()
            st.subheader("自定义出题Prompt")
            default_prompt = f"""# 药学规培处方审核出题规范
## 基础规则
1. 依据规培大纲：{st.session_state.outline_content}
2. 大纲掌握等级仅选用：{master_level}
3. 每题不适宜点必须包含难度分级：{diff_list}，推荐搭配1个L1 + 2个L2 + 1个L3
4. 每题包含3-5个不适宜用药错误点，每题至少覆盖2个人体器官系统
5. 出题场景：{scene}
6. 考核器官系统：{organ_sys}；指定药物：{select_drugs}
7. 核心考核主题：{topic_list}
8. 生成题型：{q_type}，总题量：{q_count}道
9. 每题附带【标准答案】+【详细药学解析】，标注错误点难度、掌握等级
输出格式：编号清晰，专业严谨，无多余闲聊"""
            if sel_tpl != "无":
                tpl_data = [t for t in all_tpl if t["template_name"] == sel_tpl][0]
                input_prompt = st.text_area("当前Prompt模板", value=tpl_data["prompt_content"], height=200)
            else:
                input_prompt = st.text_area("自定义出题Prompt", value=default_prompt, height=200)
            save_tpl_name = st.text_input("保存当前模板名称（留空不保存）")
            if st.button("保存Prompt模板") and save_tpl_name:
                save_prompt_template(save_tpl_name, input_prompt, model_type)
                st.success("模板已保存")
                st.rerun()
            run_btn = st.button("🤖 AI批量生成题目", type="primary", use_container_width=True)
        with col_right:
            st.subheader("生成试题预览 & 导出试卷")
            if run_btn:
                if not st.session_state.outline_content.strip():
                    st.warning("请先导入规培大纲！编辑框不能为空")
                elif not st.session_state.ai_client:
                    st.warning("侧边栏填写有效的API Key！")
                else:
                    with st.spinner("AI正在生成药学试题，请等待..."):
                        try:
                            res_text = st.session_state.ai_client.chat_completion(input_prompt)
                            st.session_state.generated_questions = res_text
                            st.success("题目生成完成，自动存入本地题库！")
                            db_data = {
                                "title": f"{scene}药学考核试题",
                                "full_content": res_text,
                                "scene": scene,
                                "organ_systems": organ_sys,
                                "topics": topic_list,
                                "diff_levels": diff_list,
                                "q_types": q_type,
                                "master_levels": master_level
                            }
                            add_question_to_db(db_data)
                        except Exception as e:
                            st.error(f"AI接口调用失败：{str(e)}")
            st.text_area("试题全文", value=st.session_state.generated_questions, height=350)
            if st.session_state.generated_questions.strip():
                exam_title = st.text_input("试卷标题", value="临床药学规培考核试卷")
                c1, c2, c3, c4 = st.columns(4)
                with c1:
                    word_full = create_exam_word(st.session_state.generated_questions, exam_title, has_answer=True)
                    st.download_button(
                        label="Word含答案",
                        data=word_full,
                        file_name=f"{exam_title}_含答案解析.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with c2:
                    word_empty = create_exam_word(st.session_state.generated_questions, exam_title, has_answer=False)
                    st.download_button(
                        label="Word无答案",
                        data=word_empty,
                        file_name=f"{exam_title}_纯试题无答案.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with c3:
                    pdf_full = create_exam_pdf(st.session_state.generated_questions, exam_title, has_answer=True)
                    st.download_button(
                        label="PDF含答案",
                        data=pdf_full,
                        file_name=f"{exam_title}_含答案解析.pdf",
                        mime="application/pdf"
                    )
                with c4:
                    pdf_empty = create_exam_pdf(st.session_state.generated_questions, exam_title, has_answer=False)
                    st.download_button(
                        label="PDF无答案",
                        data=pdf_empty,
                        file_name=f"{exam_title}_纯试题无答案.pdf",
                        mime="application/pdf"
                    )
    with tab2:
        st.subheader("历史题库多条件筛选")
        f1,f2,f3,f4,f5 = st.columns(5)
        with f1: filter_scene = st.selectbox("筛选场景", ["全部"] + ["门诊处方","住院医嘱","药历点评","用药教育","本科期末考试","规培考核"])
        with f2: filter_sys = st.selectbox("筛选器官系统", ["全部"] + ["抗感染系统","心血管系统","内分泌/糖尿病","肾脏泌尿"])
        with f3: filter_topic = st.selectbox("筛选考核主题", ["全部"] + ["药物相互作用","老年人用药禁忌","抗菌不合理使用"])
        with f4: filter_diff = st.selectbox("筛选难度", ["全部","L1","L2","L3"])
        with f5: filter_qtype = st.selectbox("筛选题型", ["全部","单选题","多选题","判断题","填空题","简答题"])
        q_list = query_questions(
            scene=filter_scene if filter_scene!="全部" else None,
            organ_system=filter_sys if filter_sys!="全部" else None,
            topic=filter_topic if filter_topic!="全部" else None,
            diff=filter_diff if filter_diff!="全部" else None,
            q_type=filter_qtype if filter_qtype!="全部" else None
        )
        st.info(f"共检索到 {len(q_list)} 道历史试题")
        selected_ids = []
        for q in q_list:
            ck = st.checkbox(f"【{q['id']}】{q['title']} | {q['scene']} | {q['create_time']}", key=f"ck_{q['id']}")
            if ck:
                selected_ids.append(q["id"])
            with st.expander("预览试题内容"):
                st.text(q["full_content"][:1000] + "......")
        st.session_state.selected_q_ids = selected_ids
        st.divider()
        if selected_ids:
            st.subheader("批量导出勾选题目")
            merge_title = st.text_input("合并试卷标题", value="题库筛选合并试卷")
            col_a, col_b, col_c = st.columns(3)
            selected_qs = get_question_by_ids(selected_ids)
            merge_text = "\n\n===== 分割线 =====\n\n".join([q["full_content"] for q in selected_qs])
            with col_a:
                merge_word = create_exam_word(merge_text, merge_title, True)
                st.download_button("批量导出Word试卷", data=merge_word, file_name=f"{merge_title}_批量试题.docx")
            with col_b:
                merge_pdf = create_exam_pdf(merge_text, merge_title, True)
                st.download_button("批量导出PDF试卷", data=merge_pdf, file_name=f"{merge_title}_批量试题.pdf")
            with col_c:
                excel_data = export_questions_excel(selected_qs)
                st.download_button("导出Excel完整题库表", data=excel_data, file_name="药学题库筛选结果.xlsx")
    with tab3:
        st.subheader("从题库选择题目在线自测、自动判分")
        test_q_list = query_questions()
        test_sel_id = st.selectbox("选择一套试题作答", [q["id"] for q in test_q_list], format_func=lambda x: f"ID{x} {[q['title'] for q in test_q_list if q['id']==x][0]}")
        target_q = [q for q in test_q_list if q["id"] == test_sel_id][0]
        st.text_area("试题内容", value=target_q["full_content"], height=400)
        st.info("复制试题到AI对话窗口作答，对照【答案】模块自行核对得分")

# 打包启动入口
def main():
    # 释放streamlit到临时目录运行
    temp_dir = tempfile.mkdtemp()
    app_script = os.path.join(temp_dir, "app.py")
    with open(app_script, "w", encoding="utf-8") as f:
        f.write("""
from launcher import run_app
if __name__ == "__main__":
    run_app()
""")
    sys.argv = [
        "streamlit", "run", app_script,
        "--server.headless=true",
        "--server.port=8501",
        "--browser.gatherUsageStats=false",
        "--server.maxUploadSize=200"
    ]
    stcli.main()
    # 清理临时文件夹
    shutil.rmtree(temp_dir, ignore_errors=True)

if __name__ == "__main__":
    main()